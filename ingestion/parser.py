"""
parser.py
Converts raw uploaded files into structured text with page metadata.
Supports: PDF (PyMuPDF), DOCX (python-docx), plain text.
"""
from __future__ import annotations

import logging
from dataclasses import dataclass, field
from pathlib import Path

logger = logging.getLogger(__name__)

SUPPORTED_MIME_TYPES = {
    "application/pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "text/plain",
    "text/markdown",
}

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md"}


@dataclass
class PageContent:
    page_number: int
    text: str
    char_offset: int = 0  # cumulative offset within the full document text


@dataclass
class ParsedDocument:
    file_path: str
    mime_type: str
    pages: list[PageContent] = field(default_factory=list)
    full_text: str = ""
    page_count: int = 0

    def __post_init__(self) -> None:
        if self.pages and not self.full_text:
            self.full_text = "\n\n".join(p.text for p in self.pages)
            self.page_count = len(self.pages)


def parse_document(file_path: str, mime_type: str) -> ParsedDocument:
    """
    Dispatch to the right parser based on mime_type.
    Returns a ParsedDocument with per-page text.
    """
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")

    suffix = path.suffix.lower()
    if mime_type and mime_type not in SUPPORTED_MIME_TYPES:
        raise ValueError(f"Unsupported file type: {mime_type} / {suffix}")

    if mime_type == "application/pdf" or suffix == ".pdf":
        return _parse_pdf(file_path)
    elif (
        mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        or suffix == ".docx"
    ):
        return _parse_docx(file_path)
    elif mime_type in ("text/plain", "text/markdown") or suffix in (".txt", ".md"):
        return _parse_text(file_path, mime_type)
    else:
        raise ValueError(f"Unsupported file type: {mime_type} / {suffix}")


def _parse_pdf(file_path: str) -> ParsedDocument:
    try:
        import fitz  # PyMuPDF
    except ImportError:
        raise ImportError("PyMuPDF is required for PDF parsing: pip install pymupdf")

    pages: list[PageContent] = []
    char_offset = 0

    with fitz.open(file_path) as doc:
        for page_num, page in enumerate(doc, start=1):
            text = page.get_text("text").strip()
            if text:
                pages.append(
                    PageContent(
                        page_number=page_num,
                        text=text,
                        char_offset=char_offset,
                    )
                )
                char_offset += len(text) + 2  # account for \n\n separator

    if not pages:
        logger.warning("PDF produced no text — may be scanned: %s", file_path)

    return ParsedDocument(file_path=file_path, mime_type="application/pdf", pages=pages)


def _parse_docx(file_path: str) -> ParsedDocument:
    try:
        from docx import Document as DocxDocument
    except ImportError:
        raise ImportError("python-docx is required: pip install python-docx")

    doc = DocxDocument(file_path)
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    full_text = "\n\n".join(paragraphs)

    # DOCX has no native pages — treat as one page
    page = PageContent(page_number=1, text=full_text, char_offset=0)
    mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    return ParsedDocument(file_path=file_path, mime_type=mime, pages=[page])


def _parse_text(file_path: str, mime_type: str) -> ParsedDocument:
    text = Path(file_path).read_text(encoding="utf-8", errors="replace").strip()
    page = PageContent(page_number=1, text=text, char_offset=0)
    return ParsedDocument(file_path=file_path, mime_type=mime_type, pages=[page])
