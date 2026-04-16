"""
test_parser.py
Unit tests for ingestion/parser.py
Tests plain text and DOCX parsing without external API calls.
PDF tests are skipped if pymupdf is not installed.
"""
import pytest
from pathlib import Path
from ingestion.parser import (
    parse_document,
    ParsedDocument,
    PageContent,
    SUPPORTED_EXTENSIONS,
)


class TestParseText:
    def test_parses_txt_file(self, sample_txt_file):
        result = parse_document(str(sample_txt_file), "text/plain")

        assert isinstance(result, ParsedDocument)
        assert result.page_count == 1
        assert len(result.pages) == 1
        assert "machine learning" in result.full_text.lower()

    def test_txt_page_number_is_one(self, sample_txt_file):
        result = parse_document(str(sample_txt_file), "text/plain")
        assert result.pages[0].page_number == 1

    def test_parses_md_file(self, tmp_path):
        md_file = tmp_path / "doc.md"
        md_file.write_text("# Title\n\nSome content here.", encoding="utf-8")

        result = parse_document(str(md_file), "text/markdown")
        assert "Title" in result.full_text

    def test_char_offset_initialised_zero(self, sample_txt_file):
        result = parse_document(str(sample_txt_file), "text/plain")
        assert result.pages[0].char_offset == 0

    def test_full_text_matches_page_text(self, sample_txt_file):
        result = parse_document(str(sample_txt_file), "text/plain")
        assert result.full_text == result.pages[0].text


class TestParseDocx:
    def test_parses_docx_file(self, tmp_path):
        pytest.importorskip("docx")
        from docx import Document

        doc = Document()
        doc.add_paragraph("First paragraph about AI.")
        doc.add_paragraph("Second paragraph about ML.")
        docx_path = tmp_path / "test.docx"
        doc.save(str(docx_path))

        result = parse_document(str(docx_path), "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        assert "First paragraph" in result.full_text
        assert "Second paragraph" in result.full_text
        assert result.page_count == 1

    def test_empty_docx_produces_empty_text(self, tmp_path):
        pytest.importorskip("docx")
        from docx import Document

        doc = Document()
        docx_path = tmp_path / "empty.docx"
        doc.save(str(docx_path))

        result = parse_document(str(docx_path), "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        assert result.page_count == 1
        assert result.pages[0].text == ""


class TestParsePdf:
    def test_parses_pdf_file(self, tmp_path):
        fitz = pytest.importorskip("fitz")

        pdf_path = tmp_path / "test.pdf"
        doc = fitz.open()
        page = doc.new_page()
        page.insert_text((72, 72), "Hello from a PDF document about natural language processing.")
        doc.save(str(pdf_path))
        doc.close()

        result = parse_document(str(pdf_path), "application/pdf")
        assert result.page_count >= 1
        assert "Hello" in result.full_text

    def test_pdf_page_numbers_start_at_one(self, tmp_path):
        fitz = pytest.importorskip("fitz")

        pdf_path = tmp_path / "multi.pdf"
        doc = fitz.open()
        for i in range(3):
            page = doc.new_page()
            page.insert_text((72, 72), f"Page content {i + 1}")
        doc.save(str(pdf_path))
        doc.close()

        result = parse_document(str(pdf_path), "application/pdf")
        page_numbers = [p.page_number for p in result.pages]
        assert page_numbers == sorted(page_numbers)
        assert page_numbers[0] == 1


class TestParseErrors:
    def test_unsupported_mime_type_raises(self, tmp_path):
        f = tmp_path / "file.xyz"
        f.write_text("data")
        with pytest.raises(ValueError, match="Unsupported"):
            parse_document(str(f), "application/octet-stream")

    def test_missing_file_raises(self):
        with pytest.raises(FileNotFoundError):
            parse_document("/nonexistent/path/file.txt", "text/plain")
